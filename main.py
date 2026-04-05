import io
import requests
import io
import re
import time
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from fastapi.responses import StreamingResponse
from fastapi import FastAPI, Depends, HTTPException, status, File, UploadFile, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import Optional
import os, shutil, uuid, json

import database
import auth

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"],
                   allow_headers=["*"])

if not os.path.exists("uploads"): os.makedirs("uploads")
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")

database.create_tables()


def get_db():
    db = database.SessionLocal()
    try:
        yield db
    finally:
        db.close()


def get_user_level(points: int):
    if points < 50:
        return "🌱 씨앗"
    elif points < 150:
        return "🌿 새싹"
    elif points < 300:
        return "🍀 잎새"
    elif points < 500:
        return "🌳 나무"
    else:
        return "👑 열매"


class NotificationManager:
    def __init__(self):
        self.connections: dict[str, WebSocket] = {}

    async def connect(self, websocket: WebSocket, username: str):
        await websocket.accept(); self.connections[username] = websocket

    def disconnect(self, username: str):
        if username in self.connections: del self.connections[username]

    async def send_personal_message(self, message: str, username: str):
        if username in self.connections: await self.connections[username].send_text(message)


notifier = NotificationManager()


@app.websocket("/ws/notify/{username}")
async def websocket_notify(websocket: WebSocket, username: str):
    await notifier.connect(websocket, username)
    try:
        while True: await websocket.receive_text()
    except WebSocketDisconnect:
        notifier.disconnect(username)


class ConnectionManager:
    def __init__(self): self.active_connections: list[WebSocket] = []

    async def connect(self, websocket: WebSocket): await websocket.accept(); self.active_connections.append(websocket)

    def disconnect(self, websocket: WebSocket): self.active_connections.remove(websocket)

    async def broadcast(self, message: str):
        for connection in self.active_connections: await connection.send_text(message)


chat_manager = ConnectionManager()


@app.websocket("/ws/chat")
async def websocket_chat(websocket: WebSocket):
    await chat_manager.connect(websocket)
    try:
        while True: data = await websocket.receive_text(); await chat_manager.broadcast(data)
    except WebSocketDisconnect:
        chat_manager.disconnect(websocket)


class UserCreate(BaseModel): username: str; nickname: str; password: str


class UserLogin(BaseModel): username: str; password: str


class PostCreate(BaseModel):
    username: str;
    password: str;
    title: str;
    content: str;
    category: str;
    file_url: Optional[str] = None
    # 💡 [새로 추가됨] 공고일 경우 들어오는 데이터
    deadline: Optional[str] = None
    external_link: Optional[str] = None


class PostUpdate(BaseModel):
    username: str;
    title: str;
    content: str;
    category: str;
    file_url: Optional[str] = None
    deadline: Optional[str] = None
    external_link: Optional[str] = None


class CommentCreate(BaseModel): username: str; password: str; post_id: int; content: str; parent_id: Optional[
    int] = None


class CommentUpdate(BaseModel): username: str; content: str


class UserUpdateInfo(BaseModel): new_nickname: Optional[str] = None; new_password: Optional[str] = None


@app.post("/upload")
def upload_file(file: UploadFile = File(...)):
    ext = file.filename.split('.')[-1];
    filename = f"{uuid.uuid4()}.{ext}";
    file_path = f"uploads/{filename}"
    with open(file_path, "wb") as buffer: shutil.copyfileobj(file.file, buffer)
    return {"file_url": f"https://announcer-project.onrender.com/{file_path}"}


@app.post("/signup")
def signup(user_data: UserCreate, db: Session = Depends(get_db)):
    existing_user = db.query(database.User).filter(database.User.username == user_data.username).first()
    if existing_user: raise HTTPException(status_code=400, detail="이미 존재하는 아이디입니다.")
    hashed_pw = auth.get_password_hash(user_data.password)
    new_user = database.User(username=user_data.username, nickname=user_data.nickname, hashed_password=hashed_pw,
                             points=0)
    db.add(new_user);
    db.commit()
    return {"message": "회원가입 완료"}


@app.post("/login")
def login(user_data: UserLogin, db: Session = Depends(get_db)):
    user = db.query(database.User).filter(database.User.username == user_data.username).first()
    if not user or not auth.verify_password(user_data.password, user.hashed_password): raise HTTPException(
        status_code=401)
    return {"message": "로그인 성공", "nickname": user.nickname, "level": get_user_level(user.points), "points": user.points}


@app.post("/posts")
def create_post(post_data: PostCreate, db: Session = Depends(get_db)):
    user = db.query(database.User).filter(database.User.username == post_data.username).first()
    if not user or not auth.verify_password(post_data.password, user.hashed_password): raise HTTPException(
        status_code=401)

    # 💡 [수정됨] 마감일과 외부 링크도 DB에 저장
    new_post = database.Post(
        title=post_data.title, content=post_data.content, category=post_data.category,
        file_url=post_data.file_url, deadline=post_data.deadline, external_link=post_data.external_link,
        user_id=user.id
    )
    db.add(new_post);
    user.points += 10;
    db.commit()
    return {"message": "등록 완료"}


# ==========================================
# 🚨 [새로 추가됨] 공채 D-Day 전광판 전용 API
# ==========================================
@app.get("/announcements", summary="공채 전광판용 데이터 조회")
def get_announcements(db: Session = Depends(get_db)):
    # 카테고리가 '공고'인 글만 모두 가져옵니다.
    announcements = db.query(database.Post).filter(database.Post.category == '공고').all()

    result = []
    for post in announcements:
        if post.deadline:
            result.append({
                "글번호": post.id, "제목": post.title,
                "마감일": post.deadline, "링크": post.external_link
            })

    # 마감일(문자열) 순서대로 오름차순 정렬 (가장 촉박한 게 앞으로 오게)
    result.sort(key=lambda x: x["마감일"])
    return result


@app.get("/posts")
def get_posts(skip: int = 0, limit: int = 5, search: Optional[str] = None, category: Optional[str] = "전체",
              sort_by: Optional[str] = "latest", db: Session = Depends(get_db)):
    query = db.query(database.Post)
    if category and category != '전체': query = query.filter(database.Post.category == category)
    if search:
        search_formatted = f"%{search}%"
        query = query.filter(database.Post.title.like(search_formatted) | database.Post.content.like(search_formatted))
    all_posts = query.all()
    if sort_by == "popular":
        all_posts.sort(key=lambda x: len(x.likes), reverse=True)
    else:
        all_posts.sort(key=lambda x: x.id, reverse=True)
    total_count = len(all_posts)
    paged_posts = all_posts[skip: skip + limit]

    result = []
    for post in paged_posts:
        result.append({
            "글번호": post.id, "제목": post.title, "내용": post.content, "카테고리": getattr(post, 'category', '자유'),
            "file_url": post.file_url, "deadline": post.deadline, "external_link": post.external_link,
            "작성자": post.author.nickname, "작성자등급": get_user_level(post.author.points),
            "작성시간": post.created_at.strftime("%Y-%m-%d %H:%M"),
            "좋아요수": len(post.likes), "좋아요누른사람들": [like.user.nickname for like in post.likes]
        })
    return {"total_count": total_count, "posts": result}


@app.get("/posts/{post_id}")
def get_post(post_id: int, db: Session = Depends(get_db)):
    post = db.query(database.Post).filter(database.Post.id == post_id).first()
    if not post: raise HTTPException(status_code=404)
    return {
        "글번호": post.id, "제목": post.title, "내용": post.content, "카테고리": getattr(post, 'category', '자유'),
        "file_url": post.file_url, "deadline": post.deadline, "external_link": post.external_link,
        "작성자": post.author.nickname, "작성자등급": get_user_level(post.author.points),
        "작성시간": post.created_at.strftime("%Y-%m-%d %H:%M"),
        "좋아요수": len(post.likes), "좋아요누른사람들": [like.user.nickname for like in post.likes]
    }


@app.put("/posts/{post_id}")
def update_post(post_id: int, post_data: PostUpdate, db: Session = Depends(get_db)):
    post = db.query(database.Post).filter(database.Post.id == post_id).first()
    if not post: raise HTTPException(status_code=404)
    if post.author.username != post_data.username: raise HTTPException(status_code=403)
    post.title = post_data.title;
    post.content = post_data.content;
    post.category = post_data.category
    if post_data.file_url: post.file_url = post_data.file_url
    if post_data.deadline: post.deadline = post_data.deadline
    if post_data.external_link: post.external_link = post_data.external_link
    db.commit()
    return {"message": "수정 완료"}


@app.delete("/posts/{post_id}")
def delete_post(post_id: int, username: str, db: Session = Depends(get_db)):
    post = db.query(database.Post).filter(database.Post.id == post_id).first()
    if not post: raise HTTPException(status_code=404)
    if post.author.username != username: raise HTTPException(status_code=403)
    post.author.points = max(0, post.author.points - 10)
    db.delete(post);
    db.commit()
    return {"message": "삭제 완료"}


@app.post("/posts/{post_id}/like")
async def toggle_like(post_id: int, username: str, db: Session = Depends(get_db)):
    user = db.query(database.User).filter(database.User.username == username).first()
    post = db.query(database.Post).filter(database.Post.id == post_id).first()
    if not user or not post: raise HTTPException(status_code=404)

    existing_like = db.query(database.PostLike).filter(database.PostLike.post_id == post_id,
                                                       database.PostLike.user_id == user.id).first()
    if existing_like:
        post.author.points = max(0, post.author.points - 3)
        db.delete(existing_like);
        db.commit()
        return {"message": "좋아요 취소"}
    else:
        db.add(database.PostLike(user_id=user.id, post_id=post_id))
        if post.author.username != username:
            post.author.points += 3
            msg = json.dumps({"text": f"❤️ {user.nickname}님이 회원님의 [{post.title}] 글을 좋아합니다. (+3점)"})
            await notifier.send_personal_message(msg, post.author.username)
        db.commit();
        return {"message": "좋아요 완료"}


@app.post("/comments")
async def create_comment(comment_data: CommentCreate, db: Session = Depends(get_db)):
    user = db.query(database.User).filter(database.User.username == comment_data.username).first()
    if not user or not auth.verify_password(comment_data.password, user.hashed_password): raise HTTPException(
        status_code=401)

    new_comment = database.Comment(content=comment_data.content, user_id=user.id, post_id=comment_data.post_id,
                                   parent_id=comment_data.parent_id)
    db.add(new_comment);
    user.points += 5;
    db.commit()

    post = db.query(database.Post).filter(database.Post.id == comment_data.post_id).first()
    if post and post.author.username != user.username:
        msg = json.dumps({"text": f"💬 {user.nickname}님이 회원님의 글에 댓글을 남겼습니다."})
        await notifier.send_personal_message(msg, post.author.username)
    return {"message": "등록 완료"}


@app.get("/posts/{post_id}/comments")
def get_comments(post_id: int, sort_by: str = "latest", db: Session = Depends(get_db)):
    comments = db.query(database.Comment).filter(database.Comment.post_id == post_id).all()
    if sort_by == "popular":
        comments.sort(key=lambda x: len(x.likes), reverse=True)
    else:
        comments.sort(key=lambda x: x.id, reverse=False)
    result = []
    for c in comments:
        result.append({
            "댓글번호": c.id, "내용": c.content, "작성자": c.author.nickname, "작성자등급": get_user_level(c.author.points),
            "작성시간": c.created_at.strftime("%Y-%m-%d %H:%M"), "좋아요수": len(c.likes),
            "좋아요누른사람들": [like.user.nickname for like in c.likes], "부모댓글번호": c.parent_id
        })
    return result


@app.put("/comments/{comment_id}")
def update_comment(comment_id: int, comment_data: CommentUpdate, db: Session = Depends(get_db)):
    comment = db.query(database.Comment).filter(database.Comment.id == comment_id).first()
    if not comment: raise HTTPException(status_code=404)
    if comment.author.username != comment_data.username: raise HTTPException(status_code=403)
    comment.content = comment_data.content;
    db.commit();
    return {"message": "수정 완료"}


@app.delete("/comments/{comment_id}")
def delete_comment(comment_id: int, username: str, db: Session = Depends(get_db)):
    comment = db.query(database.Comment).filter(database.Comment.id == comment_id).first()
    if not comment: raise HTTPException(status_code=404)
    if comment.author.username != username: raise HTTPException(status_code=403)
    comment.author.points = max(0, comment.author.points - 5)
    db.delete(comment);
    db.commit();
    return {"message": "삭제 완료"}


@app.post("/comments/{comment_id}/like")
async def toggle_comment_like(comment_id: int, username: str, db: Session = Depends(get_db)):
    user = db.query(database.User).filter(database.User.username == username).first()
    comment = db.query(database.Comment).filter(database.Comment.id == comment_id).first()
    if not user or not comment: raise HTTPException(status_code=404)
    existing = db.query(database.CommentLike).filter(database.CommentLike.comment_id == comment_id,
                                                     database.CommentLike.user_id == user.id).first()
    if existing:
        comment.author.points = max(0, comment.author.points - 2)
        db.delete(existing);
        db.commit();
        return {"message": "좋아요 취소"}
    else:
        db.add(database.CommentLike(user_id=user.id, comment_id=comment_id))
        if comment.author.username != username:
            comment.author.points += 2
            msg = json.dumps({"text": f"👍 {user.nickname}님이 회원님의 댓글을 좋아합니다. (+2점)"})
            await notifier.send_personal_message(msg, comment.author.username)
        db.commit();
        return {"message": "좋아요 완료"}


@app.get("/users/{username}/activity")
def get_user_activity(username: str, db: Session = Depends(get_db)):
    user = db.query(database.User).filter(database.User.username == username).first()
    if not user: raise HTTPException(status_code=404)
    my_posts = [{"글번호": p.id, "제목": p.title, "작성시간": p.created_at.strftime("%Y-%m-%d %H:%M")} for p in user.posts]
    my_comments = [{"댓글번호": c.id, "내용": c.content, "원문번호": c.post_id, "작성시간": c.created_at.strftime("%Y-%m-%d %H:%M")}
                   for c in user.comments]
    liked_post_records = db.query(database.PostLike).filter(database.PostLike.user_id == user.id).all()
    liked_posts = [{"글번호": rec.post.id, "제목": rec.post.title, "작성자": rec.post.author.nickname} for rec in
                   liked_post_records]
    return {"nickname": user.nickname, "points": user.points, "level": get_user_level(user.points),
            "my_posts": my_posts[::-1], "my_comments": my_comments[::-1], "liked_posts": liked_posts[::-1]}


@app.put("/users/{username}")
def update_user_info(username: str, update_data: UserUpdateInfo, db: Session = Depends(get_db)):
    user = db.query(database.User).filter(database.User.username == username).first()
    if not user: raise HTTPException(status_code=404)
    if update_data.new_nickname: user.nickname = update_data.new_nickname
    if update_data.new_password: user.hashed_password = auth.get_password_hash(update_data.new_password)
    db.commit();
    return {"message": "수정 완료", "new_nickname": user.nickname}


class GenerateRequest(BaseModel): username: str; password: str


# ==========================================
# 🤖 [복구 완료!] 네이버TV 셀레니움 실시간 크롤링 (앵커 3 + 단신 7)
# ==========================================
class GenerateRequest(BaseModel):
    username: str
    password: str


# 💡 크롤링을 도와주는 헬퍼 함수들
def parse_time_to_seconds(time_str):
    try:
        parts = time_str.split(':')
        return int(parts[0]) * 60 + int(parts[1])
    except:
        return 999


def clean_script(text):
    text = re.sub(r'#\w+', '', text)
    if "연합뉴스TV 기사문의" in text: text = text.split("연합뉴스TV 기사문의")[0]
    return text.strip()


def click_more_button(driver, wait):
    try:
        btn = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "button[class*='button_more']")))
        driver.execute_script("arguments[0].click();", btn)
        time.sleep(0.5)
    except:
        pass


def set_style(run, font_name="Malgun Gothic", size=13, bold=False, color_rgb=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color_rgb: run.font.color.rgb = color_rgb


@app.post("/generate-script", summary="MBC/연합뉴스 셀레니움 크롤링 원고 생성")
def generate_script(req: GenerateRequest, db: Session = Depends(get_db)):
    user = db.query(database.User).filter(database.User.username == req.username).first()
    if not user or not auth.verify_password(req.password, user.hashed_password):
        raise HTTPException(status_code=401, detail="인증 실패")

    # 1. 셀레니움 백그라운드(화면 없는) 실행 옵션 세팅
    options = Options()
    options.add_argument("--headless=new")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--window-size=1920,1080")
    options.add_argument(
        "user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")

    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)

    # 2. 워드 문서 뼈대 만들기
    doc = Document()
    for sec in doc.sections:
        sec.top_margin, sec.bottom_margin = Pt(40), Pt(40)
        sec.left_margin, sec.right_margin = Pt(50), Pt(50)

    try:
        # ----------------------------------------------------
        # [1] MBC 앵커멘트 수집 (딱 3개)
        # ----------------------------------------------------
        doc.add_heading('■ 앵커멘트 (총 3개 / 출처: MBC)', level=1)
        driver.get("https://tv.naver.com/imnews?tab=clip")
        time.sleep(3)

        target_links = []
        while len(target_links) < 10:  # 여유있게 링크 확보
            items = driver.find_elements(By.CSS_SELECTOR, "a.ClipCardV2_link_thumbnail__NWYf1")
            for item in items:
                if len(target_links) >= 10: break
                try:
                    sec = parse_time_to_seconds(
                        item.find_element(By.CSS_SELECTOR, "span.ClipCardV2_playtime__IHYFQ").text)
                    link = item.get_attribute("href")
                    if 140 <= sec <= 170 and link not in target_links:
                        target_links.append(link)
                except:
                    continue
            if len(target_links) >= 10: break
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(2)

        wait = WebDriverWait(driver, 5)
        success_count = 0

        for link in target_links:
            if success_count >= 3: break  # 목표치 3개 달성 시 종료
            driver.get(link)
            time.sleep(2)

            try:
                click_more_button(driver, wait)
                body = wait.until(EC.presence_of_element_located(
                    (By.CSS_SELECTOR, "div.ArticleSection_scroll_wrap__ZaUDW"))).text.strip()

                if "◀ 앵커 ▶" not in body: continue
                if "◀ 리포트 ▶" in body: body = body.split("◀ 리포트 ▶")[0]
                if "◀ 기자 ▶" in body: body = body.split("◀ 기자 ▶")[0]

                body = clean_script(re.sub(r'\[.*?\]', '', body).replace("◀ 앵커 ▶", ""))

                success_count += 1
                p = doc.add_paragraph()
                run_title = p.add_run(f"<앵커멘트 {success_count}>\n")
                set_style(run_title, size=14, bold=True, color_rgb=RGBColor(0, 112, 192))

                run_body = p.add_run(body)
                set_style(run_body, size=13)
                p.paragraph_format.line_spacing = 1.6
                p.paragraph_format.space_after = Pt(20)
            except:
                continue

        # ----------------------------------------------------
        # [2] 연합뉴스 단신 수집 (딱 7개, 무예독 없이 통합)
        # ----------------------------------------------------
        doc.add_heading('■ 단신 (총 7개 / 출처: 연합뉴스TV)', level=1)
        driver.get("https://tv.naver.com/yonhapnewstv?tab=clip")
        time.sleep(3)

        target_links = []
        while len(target_links) < 15:  # 여유있게 확보
            items = driver.find_elements(By.CSS_SELECTOR, "a.ClipCardV2_link_thumbnail__NWYf1")
            for item in items:
                if len(target_links) >= 15: break
                try:
                    title = item.get_attribute("aria-label")
                    sec = parse_time_to_seconds(
                        item.find_element(By.CSS_SELECTOR, "span.ClipCardV2_playtime__IHYFQ").text)
                    link = item.get_attribute("href")
                    # 40~53초 영상 중 속보 제외
                    if 40 <= sec <= 53 and "[속보]" not in title and link not in target_links:
                        target_links.append(link)
                except:
                    continue
            if len(target_links) >= 15: break
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(2)

        success_count = 0
        for link in target_links:
            if success_count >= 7: break  # 목표치 7개 달성 시 종료
            driver.get(link)
            try:
                click_more_button(driver, wait)
                body = clean_script(wait.until(EC.presence_of_element_located(
                    (By.CSS_SELECTOR, "div.ArticleSection_scroll_wrap__ZaUDW"))).text.strip())

                success_count += 1
                p = doc.add_paragraph()
                run_title = p.add_run(f"<단신 {success_count}>\n")
                set_style(run_title, size=14, bold=True, color_rgb=RGBColor(0, 112, 192))

                run_body = p.add_run(body)
                set_style(run_body, size=13)
                p.paragraph_format.line_spacing = 1.6
                p.paragraph_format.space_after = Pt(20)
            except:
                continue

    finally:
        # 에러가 나더라도 무조건 크롬 브라우저를 닫아줍니다 (서버 터짐 방지)
        driver.quit()

    # 3. 완성된 문서를 메모리에 저장하고 프론트엔드로 쏴주기
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Announcer_Script.docx"}
    )

if __name__ == "__main__":
    import uvicorn

    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
